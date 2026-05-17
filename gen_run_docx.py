from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = 'Microsoft YaHei'
font.size = Pt(11)

# ---- 标题 ----
title = doc.add_heading('宝可梦 MUD 游戏 — Docker 运行指南', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('只需两条命令，复制粘贴即可运行。')

# ---- 一、前提 ----
doc.add_heading('一、前提条件', level=1)
doc.add_paragraph('已安装 Docker Desktop（任务栏鲸鱼图标不转圈）。')
doc.add_paragraph('验证：打开 PowerShell，输入 docker --version，看到版本号即可。')

# ---- 二、运行 ----
doc.add_heading('二、启动游戏（2 条命令）', level=1)

doc.add_paragraph('打开 PowerShell，依次复制粘贴：')

p = doc.add_paragraph()
run = p.add_run('第 1 条：拉取镜像（只需执行一次）')
run.bold = True

doc.add_paragraph(
    'docker pull ghcr.io/erlang-yao/pokemon-mud:latest',
    style='List Bullet'
)

p = doc.add_paragraph()
run = p.add_run('第 2 条：启动游戏')
run.bold = True

doc.add_paragraph(
    'docker run -d -p 4010:4010 ghcr.io/erlang-yao/pokemon-mud:latest',
    style='List Bullet'
)

# ---- 三、打开 ----
doc.add_heading('三、开始游戏', level=1)
doc.add_paragraph('打开浏览器，输入：')
p = doc.add_paragraph()
run = p.add_run('http://localhost:4010')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)

doc.add_paragraph('看到"宝可梦世界"界面 → 输入名字 → 选御三家 → 开始冒险。')

# ---- 四、验证 ----
doc.add_heading('四、验证（可选）', level=1)
doc.add_paragraph(
    'curl http://localhost:4010/api/health',
    style='List Bullet'
)
doc.add_paragraph('返回 {"status":"ok"} 表示运行正常。')

# ---- 五、停止 ----
doc.add_heading('五、停止游戏', level=1)
doc.add_paragraph(
    'docker rm -f $(docker ps -q --filter ancestor=ghcr.io/erlang-yao/pokemon-mud:latest)',
    style='List Bullet'
)
doc.add_paragraph('或直接关掉 Docker Desktop。')

# ---- 六、常见问题 ----
doc.add_heading('六、常见问题', level=1)

problems = [
    ('docker: command not found',
     'Docker Desktop 没装或没启动。去 docker.com 下载，安装后重启电脑。'),
    ('port is already allocated',
     '4010 端口被占用。重启电脑，或在任务管理器关掉占用 4010 端口的程序。'),
    ('页面打不开',
     '确认运行了第 2 条命令。确认浏览器访问的是 http://localhost:4010（不是 5173）。'),
    ('docker pull 报 unauthorized',
     '镜像可能还在 CI 构建中，等 3-5 分钟再试。'),
]
for q, a in problems:
    p = doc.add_paragraph()
    run = p.add_run(q + '：')
    run.bold = True
    p.add_run(a)

doc.save('PokemonMUD-Guide.docx')
print('Done')
